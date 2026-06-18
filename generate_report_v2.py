import os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FLOWCHART_PATH = r'c:\\Users\\DELL\\OneDrive\\Desktop\\pdc CCP\\flowchart.png'
OUTPUT_PATH = r'c:\\Users\\DELL\\OneDrive\\Desktop\\pdc CCP\\Project_Report_v2.docx'

def set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color='C7D0D8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_padding(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def add_heading(doc, text, level=1, color=(0x1e,0x3a,0x8a)):
    h = doc.add_heading(level=level)
    h.clear()
    run = h.add_run(text)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(15) if level == 1 else Pt(13)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_body(doc, text, space_after=8):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p

def build_report():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    # Cover page – simplified for brevity
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run('PROJECT REPORT – RAG CHATBOT')
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1e,0x3a,0x8a)
    doc.add_page_break()
    # Table of contents – placeholder
    add_heading(doc, 'Table of Contents')
    for i, title in enumerate(['Executive Summary', 'System Architecture', 'Knowledge Base', 'Flowchart', 'Benchmarks', 'PDC Concepts', 'GitHub Deployment', 'Conclusion'], 1):
        p = doc.add_paragraph(f"{i}. {title}")
        p.paragraph_format.left_indent = Pt(20)
    doc.add_page_break()
    # Executive summary
    add_heading(doc, 'Executive Summary')
    add_body(doc, 'A retrieval‑augmented generation (RAG) chatbot for the Parallel & Distributed Computing course, using a parallel TF‑IDF search and Google Gemini 2.5 Flash for answer synthesis.')
    # System architecture
    add_heading(doc, 'System Architecture')
    add_body(doc, 'Three layers – Knowledge Base (markdown files), Parallel RAG Engine (ProcessPoolExecutor), Web UI (Flask + Glassmorphism).')
    # Knowledge base table (short)
    add_heading(doc, 'Knowledge Base')
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = 'File'
    hdr[1].text = 'Topics'
    for f, t in [
        ('01_intro_parallel.md', 'Flynn, parallel vs distributed'),
        ('02_laws_metrics.md', 'Amdahl, Gustafson'),
        ('03_shared_memory.md', 'OpenMP, threads'),
        ('04_distributed_mem.md', 'MPI primitives'),
        ('05_gpus_cuda.md', 'CUDA hierarchy'),
        ('06_dist_systems.md', 'CAP theorem, MapReduce')
    ]:
        row = table.add_row().cells
        row[0].text = f
        row[1].text = t
    # Flowchart section – embed image
    add_heading(doc, 'RAG Retrieval Process Flowchart')
    if os.path.exists(FLOWCHART_PATH):
        doc.add_picture(FLOWCHART_PATH, width=Inches(5.5))
        pic_para = doc.paragraphs[-1]
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figure: RAG Retrieval Flowchart')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        for r in cap.runs:
            r.font.italic = True
            r.font.size = Pt(9)
    else:
        add_body(doc, f'Flowchart image not found at {FLOWCHART_PATH}')
    # Save
    doc.save(OUTPUT_PATH)
    print(f'Report saved to {OUTPUT_PATH}')

if __name__ == '__main__':
    build_report()
